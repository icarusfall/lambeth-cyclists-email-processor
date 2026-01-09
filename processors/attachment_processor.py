"""
Attachment processor for extracting text from various file types.
Handles PDFs, Word documents, Excel files, and images.
"""

import io
from typing import List, Dict, Optional
from pathlib import Path

import pdfplumber
from docx import Document
import pandas as pd

from config.logging_config import get_logger
from models.email_data import EmailAttachment

logger = get_logger(__name__)


class AttachmentProcessor:
    """
    Processes email attachments and extracts text content.
    Supports PDFs, Word docs, Excel, and prepares images for vision API.
    """

    def __init__(self):
        """Initialize attachment processor."""
        pass

    def process_all_attachments(
        self,
        attachments: List[EmailAttachment]
    ) -> Dict[str, any]:
        """
        Process all attachments and extract text/prepare images.

        Args:
            attachments: List of EmailAttachment objects with local_path set

        Returns:
            Dictionary with:
                - combined_text: Combined text from all text-extractable attachments
                - images: List of image attachments for vision API
                - unsupported: List of unsupported attachment filenames
        """
        combined_text = []
        images = []
        unsupported = []

        for attachment in attachments:
            try:
                if not attachment.local_path:
                    logger.warning(f"Attachment {attachment.filename} has no local_path")
                    continue

                # Route to appropriate processor based on MIME type
                if attachment.mime_type == "application/pdf":
                    text = self.extract_pdf_text(attachment)
                    if text:
                        combined_text.append(f"### {attachment.filename}\n\n{text}")

                elif attachment.mime_type in [
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # .docx
                    "application/msword"  # .doc
                ]:
                    text = self.extract_word_text(attachment)
                    if text:
                        combined_text.append(f"### {attachment.filename}\n\n{text}")

                elif attachment.mime_type in [
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # .xlsx
                    "application/vnd.ms-excel",  # .xls
                    "text/csv"
                ]:
                    text = self.extract_excel_text(attachment)
                    if text:
                        combined_text.append(f"### {attachment.filename}\n\n{text}")

                elif attachment.mime_type.startswith("image/"):
                    # Images go to vision API, not text extraction
                    images.append(attachment)
                    logger.info(f"Image prepared for vision analysis: {attachment.filename}")

                else:
                    # Unsupported file type
                    unsupported.append(attachment.filename)
                    logger.info(f"Unsupported attachment type: {attachment.filename} ({attachment.mime_type})")

            except Exception as e:
                logger.error(
                    f"Error processing attachment {attachment.filename}: {e}",
                    exc_info=True
                )
                unsupported.append(attachment.filename)

        result = {
            "combined_text": "\n\n---\n\n".join(combined_text),
            "images": images,
            "unsupported": unsupported
        }

        logger.info(
            f"Processed attachments: {len(combined_text)} text, "
            f"{len(images)} images, {len(unsupported)} unsupported"
        )

        return result

    def extract_pdf_text(self, attachment: EmailAttachment) -> Optional[str]:
        """
        Extract text from PDF using pdfplumber.

        Args:
            attachment: PDF attachment

        Returns:
            Extracted text or None if extraction fails
        """
        try:
            with pdfplumber.open(attachment.local_path) as pdf:
                text_parts = []

                for i, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {i+1}]\n{page_text}")

                    # Extract tables as markdown
                    tables = page.extract_tables()
                    for j, table in enumerate(tables):
                        if table:
                            table_text = self._format_table_as_markdown(table)
                            text_parts.append(f"[Page {i+1}, Table {j+1}]\n{table_text}")

                full_text = "\n\n".join(text_parts)

                logger.info(
                    f"Extracted {len(full_text)} characters from PDF: {attachment.filename} "
                    f"({len(pdf.pages)} pages)"
                )

                return full_text if full_text.strip() else None

        except Exception as e:
            logger.error(f"Error extracting PDF text from {attachment.filename}: {e}", exc_info=True)
            return None

    def extract_word_text(self, attachment: EmailAttachment) -> Optional[str]:
        """
        Extract text from Word document (.docx).

        Args:
            attachment: Word document attachment

        Returns:
            Extracted text or None if extraction fails
        """
        try:
            doc = Document(attachment.local_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                if table_data:
                    table_text = self._format_table_as_markdown(table_data)
                    text_parts.append(f"[Table {i+1}]\n{table_text}")

            full_text = "\n\n".join(text_parts)

            logger.info(
                f"Extracted {len(full_text)} characters from Word doc: {attachment.filename}"
            )

            return full_text if full_text.strip() else None

        except Exception as e:
            logger.error(f"Error extracting Word text from {attachment.filename}: {e}", exc_info=True)
            return None

    def extract_excel_text(self, attachment: EmailAttachment) -> Optional[str]:
        """
        Extract text from Excel/CSV file.

        Args:
            attachment: Excel or CSV attachment

        Returns:
            Extracted text formatted as markdown tables or None if extraction fails
        """
        try:
            # Determine read method based on file type
            if attachment.mime_type == "text/csv" or attachment.filename.endswith(".csv"):
                df_dict = {"Sheet1": pd.read_csv(attachment.local_path)}
            else:
                # Read all sheets from Excel file
                df_dict = pd.read_excel(attachment.local_path, sheet_name=None)

            text_parts = []

            for sheet_name, df in df_dict.items():
                # Limit to first 50 rows to avoid token limits
                if len(df) > 50:
                    df = df.head(50)
                    note = f"\n(Showing first 50 of {len(df)} rows)"
                else:
                    note = ""

                # Convert to markdown table
                markdown_table = df.to_markdown(index=False)

                text_parts.append(f"[{sheet_name}]\n{markdown_table}{note}")

            full_text = "\n\n".join(text_parts)

            logger.info(
                f"Extracted {len(full_text)} characters from Excel/CSV: {attachment.filename} "
                f"({len(df_dict)} sheets)"
            )

            return full_text if full_text.strip() else None

        except Exception as e:
            logger.error(f"Error extracting Excel/CSV text from {attachment.filename}: {e}", exc_info=True)
            return None

    def _format_table_as_markdown(self, table_data: List[List[str]]) -> str:
        """
        Format table data as markdown table.

        Args:
            table_data: List of rows, each row is a list of cell values

        Returns:
            Markdown formatted table
        """
        if not table_data or not table_data[0]:
            return ""

        # Use first row as header
        header = table_data[0]
        rows = table_data[1:]

        # Build markdown table
        lines = []

        # Header row
        lines.append("| " + " | ".join(str(cell) for cell in header) + " |")

        # Separator row
        lines.append("| " + " | ".join("---" for _ in header) + " |")

        # Data rows (limit to 20 to avoid token issues)
        for row in rows[:20]:
            # Pad row if needed
            padded_row = row + [""] * (len(header) - len(row))
            lines.append("| " + " | ".join(str(cell) for cell in padded_row) + " |")

        if len(rows) > 20:
            lines.append(f"| ... ({len(rows) - 20} more rows) |")

        return "\n".join(lines)
